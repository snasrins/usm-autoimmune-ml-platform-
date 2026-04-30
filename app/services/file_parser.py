"""
File Parser Service
Handles multiple file formats for data ingestion:
- Structured: CSV, XLSX, Parquet, JSON, XML
- Unstructured: PDF, Word (DOCX), Images (PNG/JPG), TXT
"""
import pandas as pd
import numpy as np
from typing import Dict, List, Any, Optional, Union
from pathlib import Path
import hashlib
from datetime import datetime
import json
import io

# Document processing imports
try:
    import pdfplumber
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pytesseract
    from PIL import Image
    import cv2
    import numpy as np
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

try:
    from pdf2image import convert_from_path
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False


class FileParser:
    """Parse multiple file formats for lab data import"""
    
    # Structured data formats (direct to DataFrame)
    STRUCTURED_FORMATS = ['.xlsx', '.xls', '.csv', '.parquet', '.json', '.xml']
    
    # Unstructured formats (need OCR/NLP)
    UNSTRUCTURED_FORMATS = ['.pdf', '.docx', '.doc', '.txt', '.png', '.jpg', '.jpeg', '.tiff']
    
    SUPPORTED_FORMATS = STRUCTURED_FORMATS + UNSTRUCTURED_FORMATS
    
    def __init__(self, file_path: str):
        self.file_path = Path(file_path)
        self.file_type = self.file_path.suffix.lower()
        self.df: Optional[pd.DataFrame] = None
        self.raw_text: Optional[str] = None  # For unstructured data
        self.metadata: Dict[str, Any] = {}
        self.is_structured = self.file_type in self.STRUCTURED_FORMATS
        
    def validate_file(self) -> Dict[str, Any]:
        """Validate file before parsing"""
        validation = {
            "valid": True,
            "errors": [],
            "warnings": [],
            "filename": self.file_path.name,
            "file_path": str(self.file_path),
            "file_extension": self.file_type.lstrip('.'),
            "is_structured": self.is_structured
        }
        
        # Check file exists
        if not self.file_path.exists():
            validation["valid"] = False
            validation["errors"].append(f"File not found: {self.file_path}")
            return validation
        
        # Check file type
        if self.file_type not in self.SUPPORTED_FORMATS:
            validation["valid"] = False
            validation["errors"].append(
                f"Unsupported format: {self.file_type}. "
                f"Supported: {', '.join(self.SUPPORTED_FORMATS)}"
            )
            return validation
        
        # Check required libraries
        if self.file_type == '.pdf' and not PDF_AVAILABLE:
            validation["valid"] = False
            validation["errors"].append("PDF support not installed. Run: pip install pdfplumber PyPDF2")
            
        if self.file_type in ['.docx', '.doc'] and not DOCX_AVAILABLE:
            validation["valid"] = False
            validation["errors"].append("Word support not installed. Run: pip install python-docx")
            
        if self.file_type in ['.png', '.jpg', '.jpeg', '.tiff'] and not OCR_AVAILABLE:
            validation["warnings"].append("OCR not available. Install with: pip install pytesseract opencv-python")
        
        # Check file size (max 200MB for unstructured, 100MB for structured)
        file_size = self.file_path.stat().st_size
        max_size = 200 * 1024 * 1024 if not self.is_structured else 100 * 1024 * 1024
        if file_size > max_size:
            validation["valid"] = False
            validation["errors"].append(
                f"File too large: {file_size / (1024*1024):.2f}MB "
                f"(max {max_size / (1024*1024):.0f}MB)"
            )
            return validation
        
        # Calculate file hash
        validation["file_hash"] = self._calculate_hash()
        validation["file_size"] = file_size
        
        return validation
    
    def _calculate_hash(self) -> str:
        """Calculate SHA-256 hash of file"""
        sha256 = hashlib.sha256()
        with open(self.file_path, 'rb') as f:
            for chunk in iter(lambda: f.read(4096), b''):
                sha256.update(chunk)
        return sha256.hexdigest()
    
    def parse(self) -> Union[pd.DataFrame, str]:
        """
        Parse file and return DataFrame (structured) or text (unstructured)
        
        Returns:
            pd.DataFrame for structured formats
            str for unstructured formats
        """
        try:
            if self.is_structured:
                return self._parse_structured()
            else:
                return self._parse_unstructured()
                
        except Exception as e:
            raise Exception(f"Error parsing {self.file_type} file: {str(e)}")
    
    def _parse_structured(self) -> pd.DataFrame:
        """Parse structured data formats to DataFrame"""
        
        if self.file_type in ['.xlsx', '.xls']:
            self.df = pd.read_excel(self.file_path)
            
        elif self.file_type == '.csv':
            # Try different encodings
            for encoding in ['utf-8', 'latin-1', 'iso-8859-1', 'cp1252']:
                try:
                    self.df = pd.read_csv(self.file_path, encoding=encoding)
                    break
                except UnicodeDecodeError:
                    continue
            if self.df is None:
                raise Exception("Could not decode CSV with any common encoding")
                
        elif self.file_type == '.parquet':
            self.df = pd.read_parquet(self.file_path, engine='pyarrow')
            
        elif self.file_type == '.json':
            # Try JSON formats: records, columns, or raw
            with open(self.file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            if isinstance(data, list):
                self.df = pd.DataFrame(data)
            elif isinstance(data, dict):
                # Try different orientations
                try:
                    self.df = pd.DataFrame.from_dict(data, orient='index')
                except:
                    self.df = pd.DataFrame([data])
            else:
                raise Exception("Unsupported JSON structure")
                
        elif self.file_type == '.xml':
            # Try pandas XML parser (pandas 1.3+)
            try:
                self.df = pd.read_xml(self.file_path)
            except:
                # Fallback: parse with lxml and convert
                import xml.etree.ElementTree as ET
                tree = ET.parse(self.file_path)
                root = tree.getroot()
                
                # Extract data (assumes simple structure)
                data = []
                for child in root:
                    row = {elem.tag: elem.text for elem in child}
                    data.append(row)
                self.df = pd.DataFrame(data)
        
        # Store metadata
        self.metadata = {
            "format": "structured",
            "file_type": self.file_type,
            "row_count": len(self.df),
            "column_count": len(self.df.columns),
            "columns": self.df.columns.tolist(),
            "dtypes": {col: str(dtype) for col, dtype in self.df.dtypes.items()},
            "missing_values": self.df.isnull().sum().to_dict(),
            "parsed_at": datetime.now().isoformat()
        }
        
        return self.df
    
    def _parse_unstructured(self) -> str:
        """Parse unstructured formats to raw text"""
        
        if self.file_type == '.txt':
            # Plain text
            with open(self.file_path, 'r', encoding='utf-8', errors='ignore') as f:
                self.raw_text = f.read()
                
        elif self.file_type == '.pdf':
            self.raw_text = self._extract_pdf_text()
            
        elif self.file_type in ['.docx', '.doc']:
            self.raw_text = self._extract_word_text()
            
        elif self.file_type in ['.png', '.jpg', '.jpeg', '.tiff']:
            self.raw_text = self._extract_image_text_ocr()
        
        # Store metadata
        self.metadata = {
            "format": "unstructured",
            "file_type": self.file_type,
            "text_length": len(self.raw_text) if self.raw_text else 0,
            "word_count": len(self.raw_text.split()) if self.raw_text else 0,
            "parsed_at": datetime.now().isoformat(),
            "ocr_used": self.file_type in ['.png', '.jpg', '.jpeg', '.tiff', '.pdf']
        }
        
        return self.raw_text
    
    def _extract_pdf_text(self) -> str:
        """Extract text from PDF"""
        if not PDF_AVAILABLE:
            raise Exception("PDF support not installed")
        
        text_parts = []
        
        # Try pdfplumber first (better for structured PDFs)
        try:
            with pdfplumber.open(self.file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                        
            if text_parts:
                return "\n\n".join(text_parts)
        except Exception as e:
            print(f"pdfplumber failed: {e}, trying PyPDF2...")
        
        # Fallback to PyPDF2
        try:
            with open(self.file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                        
            if text_parts:
                return "\n\n".join(text_parts)
        except Exception as e:
            print(f"PyPDF2 failed: {e}")
        
        # If text extraction fails, try OCR (if available)
        if OCR_AVAILABLE and PDF2IMAGE_AVAILABLE:
            return self._ocr_pdf()
        
        return ""
    
    def _ocr_pdf(self) -> str:
        """OCR scanned PDF"""
        if not (OCR_AVAILABLE and PDF2IMAGE_AVAILABLE):
            raise Exception("OCR dependencies not installed")
        
        text_parts = []
        
        # Convert PDF pages to images
        images = convert_from_path(self.file_path, dpi=300)
        
        for i, image in enumerate(images):
            # OCR each page
            text = pytesseract.image_to_string(image, lang='eng')
            if text.strip():
                text_parts.append(f"--- Page {i+1} ---\n{text}")
        
        return "\n\n".join(text_parts)
    
    def _extract_word_text(self) -> str:
        """Extract text from Word document"""
        if not DOCX_AVAILABLE:
            raise Exception("Word support not installed")
        
        doc = Document(self.file_path)
        text_parts = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract tables
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                table_text.append(" | ".join(row_text))
            text_parts.append("\n".join(table_text))
        
        return "\n\n".join(text_parts)
    
    def _extract_image_text_ocr(self) -> str:
        """Extract text from image using OCR"""
        if not OCR_AVAILABLE:
            raise Exception("OCR support not installed")
        
        # Load image
        image = Image.open(self.file_path)
        
        # Pre-process image for better OCR
        img_array = np.array(image)
        if len(img_array.shape) == 3:
            gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
        else:
            gray = img_array
        
        # Apply thresholding
        _, thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        
        # OCR
        text = pytesseract.image_to_string(Image.fromarray(thresh), lang='eng')
        
        return text
    
    def get_preview(self, rows: int = 10) -> Dict[str, Any]:
        """Get preview of parsed data"""
        if self.df is None and self.raw_text is None:
            self.parse()
        
        if self.is_structured and self.df is not None:
            # Replace NaN/inf values with None for JSON serialization
            preview_df = self.df.head(rows).replace([float('inf'), float('-inf')], None)
            preview_df = preview_df.where(pd.notna(preview_df), None)
            
            # Convert to dict and ensure all values are JSON-serializable
            preview_records = []
            for record in preview_df.to_dict('records'):
                clean_record = {}
                for key, value in record.items():
                    if pd.isna(value) or (isinstance(value, float) and np.isinf(value)):
                        clean_record[key] = None
                    elif isinstance(value, (np.integer, np.floating)):
                        clean_record[key] = float(value) if not (np.isnan(value) or np.isinf(value)) else None
                    else:
                        clean_record[key] = value
                preview_records.append(clean_record)
            
            return {
                "format": "structured",
                "columns": self.df.columns.tolist(),
                "row_count": len(self.df),
                "column_count": len(self.df.columns),
                "preview": preview_records,
                "dtypes": self.metadata.get("dtypes", {}),
                "missing_values": self.metadata.get("missing_values", {})
            }
        else:
            # Unstructured preview
            preview_text = self.raw_text[:1000] if self.raw_text else ""
            return {
                "format": "unstructured",
                "text_length": len(self.raw_text) if self.raw_text else 0,
                "word_count": len(self.raw_text.split()) if self.raw_text else 0,
                "preview_text": preview_text,
                "preview_truncated": len(self.raw_text) > 1000 if self.raw_text else False,
                "ocr_used": self.metadata.get("ocr_used", False)
            }
    
    def get_column_stats(self) -> Dict[str, Any]:
        """Get statistics for each column - with proper NaN/Infinity handling for JSON serialization"""
        if self.df is None:
            self.parse()
        
        def safe_float(value):
            """Convert to float, replacing NaN/Infinity with None for JSON compliance"""
            if pd.isna(value) or np.isinf(value):
                return None
            return float(value)
        
        def safe_list(values):
            """Clean list values for JSON serialization"""
            result = []
            for v in values:
                if pd.isna(v) or (isinstance(v, float) and np.isinf(v)):
                    result.append(None)
                elif isinstance(v, (np.integer, np.floating)):
                    result.append(safe_float(v))
                else:
                    result.append(v)
            return result
        
        stats = {}
        for col in self.df.columns:
            col_data = self.df[col]
            
            # Calculate null percentage safely
            null_count = int(col_data.isnull().sum())
            total_count = len(col_data)
            null_percent = safe_float((null_count / total_count * 100) if total_count > 0 else 0)
            
            # Get sample values and clean them
            sample_values = safe_list(col_data.dropna().head(5).tolist())
            
            stats[col] = {
                "dtype": str(col_data.dtype),
                "null_count": null_count,
                "null_percent": null_percent,
                "unique_count": int(col_data.nunique()),
                "sample_values": sample_values
            }
            
            # Add numeric stats if applicable
            if pd.api.types.is_numeric_dtype(col_data):
                stats[col].update({
                    "min": safe_float(col_data.min()),
                    "max": safe_float(col_data.max()),
                    "mean": safe_float(col_data.mean()),
                })
        
        return stats
    
    def detect_patient_column(self) -> Optional[str]:
        """Auto-detect which column contains patient IDs"""
        if self.df is None:
            self.parse()
        
        # Common patient ID column names
        patient_id_patterns = [
            'patient', 'id', 'hospital', 'mrn', 'record', 
            'hospitalization number', 'patient_id', 'subject'
        ]
        
        for col in self.df.columns:
            col_lower = col.lower()
            for pattern in patient_id_patterns:
                if pattern in col_lower:
                    return col
        
        return None
    
    def detect_demographics(self) -> Dict[str, Optional[str]]:
        """Auto-detect demographic columns"""
        if self.df is None:
            self.parse()
        
        demo_mapping = {
            'age': ['age', 'age_years', 'yr'],
            'gender': ['gender', 'sex', 'male', 'female'],
            'ethnicity': ['ethnicity', 'race', 'ethnic'],
        }
        
        detected = {}
        for demo_field, patterns in demo_mapping.items():
            detected[demo_field] = None
            for col in self.df.columns:
                col_lower = col.lower()
                for pattern in patterns:
                    if pattern in col_lower:
                        detected[demo_field] = col
                        break
                if detected[demo_field]:
                    break
        
        return detected
    
    def get_metadata(self) -> Dict[str, Any]:
        """Get file metadata"""
        return {
            **self.metadata,
            "file_path": str(self.file_path),
            "file_type": self.file_type,
            "file_name": self.file_path.name,
        }


# Usage example:
if __name__ == "__main__":
    # Test with SLE dataset
    parser = FileParser("C:/Users/Syarifah/Downloads/Dataset-Extracted/Dataset/AAM-SLE-E (real data).xlsx")
    
    # Validate
    validation = parser.validate_file()
    print("Validation:", validation)
    
    # Parse
    df = parser.parse()
    print(f"\nParsed: {len(df)} rows, {len(df.columns)} columns")
    
    # Get preview
    preview = parser.get_preview(3)
    print(f"\nColumns: {preview['columns'][:10]}")
    
    # Detect patient ID
    patient_col = parser.detect_patient_column()
    print(f"\nPatient ID column: {patient_col}")
    
    # Detect demographics
    demo = parser.detect_demographics()
    print(f"Demographics: {demo}")
